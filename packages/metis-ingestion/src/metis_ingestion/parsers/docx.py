"""DOCX text extraction via python-docx: paragraphs as blocks, table rows as TSV lines."""

from __future__ import annotations

import io

from docx import Document

from metis_ingestion._text import normalize_blocks


def extract(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
    for table in document.tables:
        rows = ["\t".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        if any(row.strip() for row in rows):
            blocks.append("\n".join(rows))
    return normalize_blocks("\n\n".join(blocks))
